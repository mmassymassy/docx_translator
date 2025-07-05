from docx import Document
from googletrans import Translator
import time

# File paths
INPUT_FILE = "origin.docx"
OUTPUT_FILE = "translated_output.docx"

# Initialize translator
translator = Translator()

# Load original document
doc = Document(INPUT_FILE)

# Translate paragraphs in-place
for para in doc.paragraphs:
    if para.text.strip():
        try:
            translated = translator.translate(para.text, src='fr', dest='en').text
            para.text = translated
            time.sleep(0.3)  # polite delay for free API
        except Exception as e:
            print(f"Error translating paragraph: {para.text[:60]}... – {e}")
            para.text = "#error#"

# Translate text in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                try:
                    translated = translator.translate(cell.text, src='fr', dest='en').text
                    cell.text = translated
                    time.sleep(0.3)
                except Exception as e:
                    print(f"Error translating table cell: {cell.text[:60]}... – {e}")
                    cell.text = "#error#"

# Save translated version
doc.save(OUTPUT_FILE)
print(f"✅ Translated document saved as: {OUTPUT_FILE}")
