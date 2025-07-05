import os
import time
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

from docx import Document
from odf.opendocument import load as load_odt
from odf.text import P
from odf import text as odf_text

from googletrans import Translator

class TranslatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("DOC/DOCX/ODT Translator - FR to EN")
        self.root.geometry("680x520")
        self.translator = Translator()
        self.file_path = ""
        self.errors = []

        self.create_widgets()

    def create_widgets(self):
        tk.Label(self.root, text="Select a DOC, DOCX or ODT file to translate:", font=("Arial", 12)).pack(pady=10)

        tk.Button(self.root, text="Browse File", font=("Arial", 10), command=self.browse_file).pack()

        self.progress = ttk.Progressbar(self.root, orient="horizontal", length=500, mode="determinate")
        self.progress.pack(pady=10)

        self.eta_label = tk.Label(self.root, text="ETA: --:--", font=("Arial", 10))
        self.eta_label.pack()

        self.log = tk.Text(self.root, height=12, width=85, font=("Consolas", 9))
        self.log.pack(pady=10)
        self.log.insert(tk.END, "Waiting for file...\n")

        self.translate_btn = tk.Button(self.root, text="Translate", font=("Arial", 10), command=self.start_translation)
        self.translate_btn.pack(pady=5)

        tk.Label(self.root, text="Made by AHMANN Massi | YouTube: @maxim_ah | Instagram: @maxim_ah",
                 font=("Arial", 9), fg="gray").pack(pady=10)

    def browse_file(self):
        filetypes = [("Word & ODT files", "*.docx *.doc *.odt")]
        self.file_path = filedialog.askopenfilename(filetypes=filetypes)
        if self.file_path:
            self.log.insert(tk.END, f"Loaded file: {self.file_path}\n")

    def start_translation(self):
        if not self.file_path:
            messagebox.showerror("Error", "No file selected.")
            return

        self.translate_btn.config(state='disabled')
        threading.Thread(target=self.translate_file).start()

    def update_progress(self, count, total, start_time):
        percent = (count / total) * 100
        self.progress["value"] = percent
        elapsed = time.time() - start_time
        remaining = (elapsed / count) * (total - count) if count else 0
        eta = f"ETA: {int(remaining//60):02d}:{int(remaining%60):02d}"
        self.eta_label.config(text=eta)

    def translate_text(self, text):
        try:
            return self.translator.translate(text, src='fr', dest='en').text
        except Exception as e:
            self.errors.append(f"{text[:40]}... → {e}")
            return "#error#"

    def translate_file(self):
        ext = os.path.splitext(self.file_path)[1].lower()
        start_time = time.time()
        self.errors.clear()

        if ext == ".docx":
            doc = Document(self.file_path)
            elements = []

            for para in doc.paragraphs:
                elements.append(para)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        elements.extend(cell.paragraphs)

            total = len(elements)
            for i, para in enumerate(elements):
                for run in para.runs:
                    if run.text.strip():
                        run.text = self.translate_text(run.text)
                self.update_progress(i + 1, total, start_time)

            # Headers and footers
            for section in doc.sections:
                for container in [section.header, section.footer]:
                    for para in container.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                run.text = self.translate_text(run.text)

            output = self.file_path.replace(".", "_translated.")
            doc.save(output)
            self.log.insert(tk.END, f"✅ Translated .docx saved: {output}\n")

        elif ext == ".odt":
            textdoc = load_odt(self.file_path)
            paras = textdoc.getElementsByType(P)
            total = len(paras)

            for i, para in enumerate(paras):
                new_text = ""
                for node in para.childNodes:
                    if node.nodeType == node.TEXT_NODE:
                        original = str(node.data).strip()
                        if original:
                            node.data = self.translate_text(original)
                self.update_progress(i + 1, total, start_time)

            output = self.file_path.replace(".", "_translated.")
            textdoc.save(output)
            self.log.insert(tk.END, f"✅ Translated .odt saved: {output}\n")

        elif ext == ".doc":
            self.log.insert(tk.END, "⚠️ .doc is not fully supported (old binary format). Please convert to .docx.\n")
        else:
            self.log.insert(tk.END, "❌ Unsupported file format.\n")

        if self.errors:
            self.log.insert(tk.END, f"\n⚠️ {len(self.errors)} translation errors occurred. Look for '#error#' in the file.\n")

        self.translate_btn.config(state='normal')

if __name__ == "__main__":
    root = tk.Tk()
    app = TranslatorApp(root)
    root.mainloop()
