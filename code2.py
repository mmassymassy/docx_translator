from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from googletrans import Translator
import os
import time

INPUT_FILE = "origin.docx"
OUTPUT_FILE = "origin_translated2.docx"

translator = Translator()

def translate_run(run):
    try:
        translated = translator.translate(run.text, src="fr", dest="en").text
        return translated
    except Exception as e:
        print(f"Error: {e} – '{run.text[:40]}...'")
        return "#error#"

def translate_paragraph(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            run.text = translate_run(run)
    return paragraph

def translate_paragraphs(paragraphs):
    for paragraph in paragraphs:
        if paragraph.text.strip():
            translate_paragraph(paragraph)

def translate_table(table):
    for row in table.rows:
        for cell in row.cells:
            translate_paragraphs(cell.paragraphs)
            for tbl in cell.tables:
                translate_table(tbl)

def translate_header_footer(header_or_footer):
    for paragraph in header_or_footer.paragraphs:
        translate_paragraph(paragraph)

def translate_docx(input_path, output_path):
    doc = Document(input_path)

    # Body
    for para in doc.paragraphs:
        translate_paragraph(para)

    # Tables
    for table in doc.tables:
        translate_table(table)

    # Headers & Footers
    for section in doc.sections:
        translate_header_footer(section.header)
        translate_header_footer(section.footer)

    doc.save(output_path)
    print(f"\n✅ Translation complete. Saved as: {output_path}")

if __name__ == "__main__":
    start = time.time()
    translate_docx(INPUT_FILE, OUTPUT_FILE)
    print(f"⏱ Done in {round(time.time() - start, 2)} seconds.")
